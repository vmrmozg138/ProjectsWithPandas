from openpyxl import load_workbook
from os import listdir
from os.path import isfile, join
from tkinter import filedialog as tkd
from docx import Document
import re
import string
import pandas as pd

#import fromword

def getTables(filename):
    doc = Document(filename)
    fullTablelist = []
    goodpattern = re.compile(r"^[^/]")
    if len(doc.tables) == 0:
        return "no tables in current document"
    else:
        for tb in doc.tables:
            onetable = []
            for r in tb.rows:
                rdata = [cell.text for cell in r.cells if (cell.text != "" and goodpattern.match(cell.text))]
                onetable.append(rdata)
            if len(onetable) > 1:
                fullTablelist.append(onetable)
        return fullTablelist

def makequestions(initialcoll):
    singlecodepattern = re.compile(r"^[0-9][0-9 ]*$")
    code_with_label_pattern = re.compile(r"^[0-9]+\.\s[А-Я]")
    qcollection = []
    for titem in initialcoll:
        quest = []
        for litem in titem:
            elem = {}
            hasname = False
            if len(litem)>1:
                checkitem = litem[0].strip()
                if len(checkitem)>1 and (checkitem[0] in string.ascii_uppercase or checkitem[0] == 'С') and checkitem[-1].isdigit() and not hasname:
                #if len(litem[0])>1 and not hasname:
                    #print(litem[0],'   ',litem[1])
                    elem.update({'qname':str(checkitem.replace('.','_').replace('С','C'))})
                    elem.update({'qlabel':str(litem[1])})
                    quest.append(elem)
                    hasname = True
            for x in litem:
                if singlecodepattern.match(x):
                    if 'code' in elem:
                        pass
                    else:
                        elem.update({'code':x.split(' ')[-1]})
                elif code_with_label_pattern.match(x):
                    elem.update(dict(zip(['code','label'],x.split('. ')[:2])))
                else:
                    if 'label' in elem:
                        pass
                    else:
                        elem.update({'label':x.strip()})
            if 'code' in elem and 'label' in elem:
                quest.append(elem)
        #print(quest, '\n')
        qcollection.append(quest)
    return qcollection

def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        return False

def getnumericcode(inputvalue):
    if not is_number(inputvalue):
        if '_' in str(inputvalue):
            result = int(inputvalue.split('_')[-1])
    else:
        result = inputvalue
    return result

def validatelabel(s):
    if not isinstance(s, str):
        result = s
    else:
        variants = [['<Банк>','<БАНК>','<БАНК ОПРОСА>'],['<Банка>','<БАНКА>','<БАНКА ОПРОСА>'],['<Банке>','<БАНКЕ>','<БАНКЕ ОПРОСА>'],['<Банку>','<БАНКУ>','<БАНКУ ОПРОСА>'],['<Банком>','<БАНКОМ>','<БАНКОМ ОПРОСА>']]
        substitutions = ['Qbank','B1_insert','B1u_insert','em1_insert','QAtributs_insert']
        for v_item in variants:
            for v_variant in v_item:
                if str(v_variant) in s:
                    #print('before: ',s)
                    ins = '{#' + substitutions[variants.index(v_item)] + '.response.label}'
                    s=s.replace(v_variant,ins)
                    #print('after:', s)
        numquotes = s.count('\"')
        if numquotes == 0 or (numquotes % 2) == 1:
            result = s
        else:
            tmps = ""
            for i in range(len(s)):
                if s[i] == "\"":
                    if (numquotes % 2) == 0 :
                        replacer = "«"
                    else:
                        replacer = "»"
                    tmps += replacer
                    numquotes-=1
                elif s[i] == "“" or s[i] == "”":
                    if s[i] == "“":
                        replacer = "«"
                    else:
                        replacer = "»"
                    tmps += replacer
                else:
                    tmps += s[i]                
            result = tmps
    return result

file_path_string = tkd.askopenfilename()
fdir = str(file_path_string[:file_path_string.rfind('/')+1])
onlyfiles = [f for f in listdir(fdir) if isfile(join(fdir, f))]
docfiles = [d for d in onlyfiles if (str(d).endswith('docx') and not str(d).startswith('~'))]
print(onlyfiles,'\n',docfiles)
doctables = getTables(fdir+docfiles[0])
all_banks_dict = makequestions(doctables)[0]
banks_dict = []

for lt_item in doctables[-1]:
    basicbankname = str(lt_item[0]).lower()
    if len(basicbankname.split(' '))>3:
        basicbankname = str(basicbankname.split(' ')[0]).replace(',','')
    if 'псб' in basicbankname:
        basicbankname = 'псб'
    if 'втб' in basicbankname:
        basicbankname = 'втб'
    for iquest in all_banks_dict:
        if basicbankname in str(iquest['label']).lower():
            banks_dict.append(iquest)

#print(all_banks_dict,'\n')
storage = []
validAlphabet = 'йцукенгшщзхъфывапролджэячсмитьбюё0123456789'

wb = load_workbook(file_path_string)
b2lists = []
for sheetitem in wb.sheetnames:
    if 'B2' in str(sheetitem):
        b2lists.append(sheetitem)

print(banks_dict)

if len(b2lists) == 0:
    print('please check lists with B1 names')
    exit()
for b2_litem in b2lists:
    subname = b2_litem.split('_')[-1]    
    worksheet = wb[b2_litem]
    lines = []
    for row in worksheet.rows:    
        lines.append([str(cell.value).strip() if str(cell.value)!='None' else '' for cell in row])
    storage.extend(lines)
    df = pd.DataFrame(storage)

    groupIndsdf = df.index[df[1].str.contains('\.')].tolist()   #лист с индексами групп(заголовков)
    b2aIndsdf = df.index[df[1].str.contains('B2a') | df[2].str.contains('B2a')].tolist()    #лист с индексами всех элементов b2a
    dfb2a = df.loc[b2aIndsdf[0]:b2aIndsdf[-1]]   #df который содержит только элементы b2a
   
    #print(dfb2a)
    #print(groupIndsdf, b2aIndsdf)
    #print(df[[1,2]])

    #add index column to each bank dict
    for bank in banks_dict:
        meaningName = re.sub("-?\s?банк", "", bank['label'].lower().strip())
        if 'псб' in meaningName:
            meaningName = 'псб' 
        print(bank['label'],' - ',meaningName)
        for ind, item in enumerate(list(df.iloc[0])):
            print(meaningName, '      ', item.strip().lower())
            if meaningName in item.strip().lower():
                bank['curListCol'] = ind
                print(ind, ', ', item)
    print('\n\n')
    metInsQuestions = ''
    b2Structure = ''
    routingSetVals = ''
    banksRoutingFilter = ''

    for bank in banks_dict:
        print('\n')
        for k,v in bank.items():
            print(k,' ',v)

    for pos in range(len(groupIndsdf)):
        startInd = int(groupIndsdf[pos])
        if pos<len(groupIndsdf)-1:            
            endInd = int(groupIndsdf[pos+1])    
        else:
            endInd = int(b2aIndsdf[0])
        dfgroup = df.loc[startInd:endInd-1]        
        dfGroupCode = dfgroup.loc[dfgroup.index[0],1].replace('В','B')
        dfgWOheader = dfgroup.loc[startInd+1:]      
        print(dfGroupCode)

        #below is making metadata inserts and routing setting values        
        for i,r in dfgroup.iterrows():
            qname = str(r[1]).replace('.','_',1).replace('.','0')+'_insert'
            routingSetVals += '\t'+qname+'.response.value = Qbank.response.value\n'
            defLabel = validatelabel(str(r[3]))
            oneFullAttr = '\t{\n\t'+ ',\n\t'.join(['\t_'+str(bank['code'])+' \"'+ ( validatelabel(r[bank['curListCol']]) if len(r[bank['curListCol']])>0 else defLabel) + '\"' for bank in banks_dict ])+'\n\t};\n\n'
            metInsQuestions += '\t'+qname+' \"\"\n\tcategorical[1..1]\n'+oneFullAttr              
            print(endInd,i)
            if '.' in str(r[1]):
                b2Structure += '\t_'+qname.split('_')[1] + ' \"{#'+qname+'.response.label}\"\n\t{\n'
            else:
                b2Structure += '\t\t_'+str(r[1]).split('_')[-1] + ' \"{#'+qname+'.response.label}\"' + (',\n' if i != endInd-1 else '\n\t}\n')    

        #now iterating over banks to create filters
        for bank in banks_dict:
            filterName = dfGroupCode.replace('.','_',1).replace('.','0')+'_filter'            
            filterList = list(dfgWOheader.loc[dfgroup[bank['curListCol']] != 'NONE',1])
            stringFilter = '{' + ','.join(['_'+str(attr).split('_')[-1] for attr in filterList]) +'}'
            bank[filterName] = stringFilter

    #making bank asking
    banksAsking = ''
    for i, r in df.iterrows():
        if any(str(r[1]).startswith(x) for x in ['B2.','В2.']):
            catName = (str(r[1]).replace('.','_',1).replace('.','0')).split('_')[1]
            banksList = '{' + ','.join(['_'+str(bank['code']) for bank in banks_dict if r[bank['curListCol']] != 'NONE']) + '}'
            banksAsking += '\tQbank*'+banksList+'\n\t\t'+'_'+catName+'.Ask()\n'

    #setting bank routing filters B2 together with B2a
    for bank in banks_dict:
        b2a_StrFilter = '{'+ ','.join(['_'+str(attr).split('_')[-1] for attr in list(dfb2a.loc[dfb2a[bank['curListCol']] != 'NONE',1])]) + '}'
        bank['b2an_filter'] = b2a_StrFilter 
        banksRoutingFilter += '\tcase{_'+bank['code']+'}\n'
        for k,v in bank.items():
            if 'filter' in k:
                banksRoutingFilter += '\t\t'+str(k).lower() + ' = ' +str(v)+'\n'

    #separate code for b2a - metadata and some routing
    b2aroutingSetVals = ''
    b2ametInsQuestions = ''
    b2aStructure = ''
    for i, r in dfb2a.iterrows():
        b2aqname = str(r[1]).replace('.','_',1).replace('.','0')+'_insert'
        b2aroutingSetVals += '\t'+qname+'.response.value = Qbank.response.value\n'
        b2adefLabel = validatelabel(str(r[3]))
        b2aoneFullAttr = '\t{\n\t'+ ',\n\t'.join(['\t_'+str(bank['code'])+' \"'+ (validatelabel(r[bank['curListCol']]) if len(r[bank['curListCol']])>0 else defLabel) + '\"' for bank in banks_dict ])+'\n\t};\n\n'
        b2ametInsQuestions += '\t'+b2aqname+' \"\"\n\tcategorical[1..1]\n'+b2aoneFullAttr   
        b2aStructure += '\t_'+str(r[1]).split('_')[-1] + ' \"{#'+b2aqname+'.response.label}\",\n'      

#below writing in 2 separate documents
    with open(fdir+subname+"_b2.txt", "w") as f:
        f.write('\t\tMETADATA:\n\n')
        f.write(metInsQuestions + '\n\n' + b2Structure)
        f.write('\n\n\t\tROUTING:\n\n')        
        f.write(routingSetVals + '\n\n' + banksRoutingFilter + '\n\n' +  banksAsking)

    with open(fdir+subname+"_b2an.txt", "w") as f1:
        f1.write('\t\tMETADATA:\n\n')
        f1.write(b2ametInsQuestions + '\n\n' + b2aStructure)
        f1.write('\n\n\t\tROUTING:\n\n')
        f1.write(b2aroutingSetVals)